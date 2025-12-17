import customtkinter as ctk
import os
import subprocess
import shutil
from tkinter import filedialog, messagebox
import transcriber
import threading
import pygame
import time
import re
from tkinterdnd2 import TkinterDnD, DND_FILES
from PIL import Image
import webbrowser
import csv
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Nota: Instala 'python-docx' para exportar a Word.")
import platform
import sys
import tkinter as tk

def resource_path(relative_path):
    """Obtiene la ruta absoluta al recurso, funcione en dev o en PyInstaller"""
    try:
        # PyInstaller crea una carpeta temporal en _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Configuración inicial
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class OpenTranscribeApp(ctk.CTk, TkinterDnD.DnDWrapper):
    def __init__(self):
        super().__init__()

        self.TkdndVersion = TkinterDnD._require(self)
        pygame.mixer.init()

        self.title("OpenTranscribe v2.0")
        self.geometry("750x700")

        try:
            # Buscamos icon.png usando la función segura
            icon_file = resource_path("icon.png")

            if os.path.exists(icon_file):
                # Para Linux/macOS (y Windows modernos con PNG) se usa iconphoto
                img_icon = tk.PhotoImage(file=icon_file)
                self.iconphoto(True, img_icon) # True aplica el icono a todas las ventanas futuras
            else:
                print(f"Advertencia: No se encontró {icon_file}")
        except Exception as e:
            print(f"Error cargando icono: {e}")

        # Configuración principal
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)

        self.unsaved_changes = False
        self.selected_file_path = ""
        self.is_playing = False
        self.total_duration = 0
        self.current_offset = 0
        self.transcript_segments = []

        self.queue_files = []
        self.is_batch_mode = False

        # ============================================================
        # 1. TÍTULO
        # ============================================================
        self.lbl_title = ctk.CTkLabel(self, text="OpenTranscribe", font=("Roboto Medium", 26))
        self.lbl_title.pack(pady=(20, 10))

        # ============================================================
        # 2. TARJETA DE CONTROL (Card UI)
        # ============================================================
        self.card_frame = ctk.CTkFrame(self, fg_color="#2b2b2b", corner_radius=15)
        self.card_frame.pack(fill="x", padx=25, pady=(0, 15))

        # Configuración de columnas de la tarjeta
        self.card_frame.grid_columnconfigure(0, weight=0) # Columna 0 ajustada al botón (NO se estira)
        self.card_frame.grid_columnconfigure(1, weight=1) # Columna 1 se estira (para el texto del archivo)
        self.card_frame.grid_columnconfigure(2, weight=0)

        # --- Fila A: Selección de Archivo ---
        # width=160 (Fijo), height=32 (Igual a guardar), SIN sticky
        self.btn_browse = ctk.CTkButton(self.card_frame, text="📂 Seleccionar Audio", command=self.select_file, width=160, height=32)
        self.btn_browse.grid(row=0, column=0, padx=(20, 10), pady=(20, 10))

        self.lbl_filename = ctk.CTkLabel(self.card_frame, text="Arrastra un archivo aquí...", text_color="gray", anchor="w")
        self.lbl_filename.grid(row=0, column=1, columnspan=2, padx=(0, 20), pady=(20, 10), sticky="ew")

        # --- Fila B: Reproductor ---
        self.frame_player = ctk.CTkFrame(self.card_frame, fg_color="transparent")
        self.frame_player.grid(row=1, column=0, columnspan=3, sticky="ew", padx=10, pady=5)

        self.btn_play = ctk.CTkButton(self.frame_player, text="▶", state="disabled", width=50, command=self.toggle_audio, fg_color="#444", height=30)
        self.btn_play.pack(side="left", padx=(5, 5))

        self.btn_stop = ctk.CTkButton(self.frame_player, text="⏹", state="disabled", width=40, command=self.stop_audio, fg_color="#800000", height=30)
        self.btn_stop.pack(side="left", padx=(0, 15))

        self.slider_audio = ctk.CTkSlider(self.frame_player, from_=0, to=1, command=self.seek_audio, height=18)
        self.slider_audio.set(0)
        self.slider_audio.pack(side="left", fill="x", expand=True, padx=(0, 10))

        self.lbl_audio_time = ctk.CTkLabel(self.frame_player, text="00:00 / 00:00", font=("Arial", 11), text_color="#aaa")
        self.lbl_audio_time.pack(side="right", padx=(0, 5))

        # --- Fila C: Configuración ---
        self.frame_settings = ctk.CTkFrame(self.card_frame, fg_color="transparent")
        self.frame_settings.grid(row=2, column=0, columnspan=3, sticky="ew", padx=15, pady=(15, 10))

        # Modelo
        ctk.CTkLabel(self.frame_settings, text="Modelo IA:", font=("Arial", 12, "bold")).pack(side="left", padx=(0, 10))
        self.combo_models = ctk.CTkComboBox(self.frame_settings, width=180, values=["Tiny (Muy rápido)", "Base (Equilibrado)", "Small (Preciso)", "Medium (Muy preciso)", "Large (Lento/Pro)"])
        self.combo_models.set("Base (Equilibrado)")
        self.combo_models.pack(side="left")

        # Switches (Usamos un frame a la derecha para apilarlos o ponerlos juntos)
        self.frame_switches = ctk.CTkFrame(self.frame_settings, fg_color="transparent")
        self.frame_switches.pack(side="right")

        self.switch_diarize = ctk.CTkSwitch(self.frame_switches, text="Detectar Hablantes 👥") # <--- NUEVO
        self.switch_diarize.pack(side="left", padx=(0, 15))

        self.switch_srt = ctk.CTkSwitch(self.frame_switches, text="Modo Subtítulos")
        self.switch_srt.pack(side="left")

        # --- Fila D: BOTONES DE ACCIÓN (Centrados y Tamaño Fijo) ---
        self.frame_big_btns = ctk.CTkFrame(self.card_frame, fg_color="transparent")
        self.frame_big_btns.grid(row=3, column=0, columnspan=3, sticky="ew", padx=15, pady=(15, 20))

        # Usamos columnas con peso para que los botones se centren en su mitad, pero no se estiren
        self.frame_big_btns.grid_columnconfigure(0, weight=1)
        self.frame_big_btns.grid_columnconfigure(1, weight=1)

        # Botón Transcribir: width=160, height=32. Quitamos sticky="ew" para que no se estire.
        self.btn_process = ctk.CTkButton(self.frame_big_btns, text="TRANSCRIBIR", font=("Arial", 12, "bold"), state="disabled", fg_color="green", hover_color="#006400", command=self.start_transcription, width=160, height=32)
        self.btn_process.grid(row=0, column=0, padx=10) # Centrado en su columna

        # Botón Cancelar: width=160, height=32.
        self.btn_cancel = ctk.CTkButton(self.frame_big_btns, text="CANCELAR", font=("Arial", 12, "bold"), state="disabled", fg_color="#8b0000", hover_color="#500000", command=self.cancel_transcription, width=160, height=32)
        self.btn_cancel.grid(row=0, column=1, padx=10) # Centrado en su columna

        # --- Fila E: Barra de Progreso ---
        self.progress_bar = ctk.CTkProgressBar(self.card_frame, height=6, corner_radius=0)
        self.progress_bar.grid(row=4, column=0, columnspan=3, sticky="ew", padx=15, pady=(0, 15))
        self.progress_bar.set(0)

        self.lbl_progress_percent = ctk.CTkLabel(self.card_frame, text="0%", font=("Arial", 10))
        self.lbl_progress_percent.place(relx=0.95, rely=0.92, anchor="e")

        # ============================================================
        # 3. ÁREA DE TEXTO
        # ============================================================
        self.frame_text = ctk.CTkFrame(self, fg_color="transparent")
        self.frame_text.pack(fill="both", expand=True, padx=25, pady=(0, 10))

        self.textbox = ctk.CTkTextbox(self.frame_text, corner_radius=10, font=("Roboto", 14), fg_color="#1d1d1d", border_width=1, border_color="#333")
        self.textbox.pack(fill="both", expand=True)
        self.textbox.insert("0.0", "El texto transcrito aparecerá aquí...\n")
        self.textbox._textbox.bind("<KeyRelease>", self.mark_as_modified)
        self.textbox._textbox.tag_config("highlight", background="#005f73", foreground="#ffffff")

        # ============================================================
        # 4. BARRA INFERIOR
        # ============================================================
        self.frame_bottom = ctk.CTkFrame(self, fg_color="transparent")
        self.frame_bottom.pack(fill="x", padx=25, pady=(0, 20))

        # Izquierda
        self.btn_help = ctk.CTkButton(self.frame_bottom, text="? Ayuda", command=self.abrir_ayuda, fg_color="#333", width=80, height=28)
        self.btn_help.pack(side="left", padx=(0, 10))

        self.btn_find = ctk.CTkButton(self.frame_bottom, text="🔍 Buscar", command=self.open_find_replace_dialog, fg_color="#333", width=80, height=28)
        self.btn_find.pack(side="left", padx=(0, 10))

        # Centro
        self.btn_sync = ctk.CTkButton(self.frame_bottom, text="🔄 Sincronizar Tiempos", command=self.sync_timestamps_manual, fg_color="#4a4a4a", width=140, height=28)
        self.btn_sync.pack(side="left")

        # Derecha (Referencia de tamaño)
        self.btn_save = ctk.CTkButton(self.frame_bottom, text="💾 Guardar Texto", command=self.save_text, fg_color="#1f6aa5", height=32, font=("Arial", 12, "bold"))
        self.btn_save.pack(side="right")

        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.drop_target_register(DND_FILES)
        self.dnd_bind('<<Drop>>', self.al_soltar_archivo)

        self.after(500, self.check_system_requirements)

    # ==========================================================
    # LÓGICA DE LA APLICACIÓN (SIN CAMBIOS FUNCIONALES)
    # ==========================================================

    def mostrar_confirmacion_oscura(self, titulo, mensaje):
        # Crear ventana emergente oscura
        dialog = ctk.CTkToplevel(self)
        dialog.title(titulo)
        dialog.geometry("350x180")
        dialog.resizable(False, False)

        # Hacemos que flote siempre encima
        dialog.attributes("-topmost", True)

        # Etiqueta con el mensaje
        lbl = ctk.CTkLabel(dialog, text=mensaje, font=("Roboto", 14), wraplength=300)
        lbl.pack(pady=30, padx=20)

        # Frame para los botones
        frame_btns = ctk.CTkFrame(dialog, fg_color="transparent")
        frame_btns.pack(pady=10)

        # Variable para guardar la respuesta
        self.respuesta_usuario = False

        def on_yes():
            self.respuesta_usuario = True
            dialog.destroy()

        def on_no():
            self.respuesta_usuario = False
            dialog.destroy()

        # Botones Sí/No estilizados
        btn_yes = ctk.CTkButton(frame_btns, text="Sí", command=on_yes, fg_color="#8b0000", hover_color="#500000", width=80)
        btn_yes.pack(side="left", padx=10)

        btn_no = ctk.CTkButton(frame_btns, text="No", command=on_no, fg_color="gray", hover_color="#555", width=80)
        btn_no.pack(side="right", padx=10)

        # Bloquear la ventana principal hasta que se responda
        dialog.transient(self)
        dialog.grab_set()
        self.wait_window(dialog)

        return self.respuesta_usuario

    def cancel_transcription(self):
        # Usamos nuestra nueva alerta oscura
        confirm = self.mostrar_confirmacion_oscura("Cancelar", "¿Seguro que quieres detener la transcripción?")

        if confirm:
            transcriber.stop_transcription()
            self.btn_cancel.configure(state="disabled")
            self.btn_process.configure(state="normal", text="Transcribir")
            self.title("OpenTranscribe (Cancelado)")
            self.progress_bar.set(0)
            self.lbl_progress_percent.configure(text="0%")

    def start_transcription(self):
        # 1. Configuración común
        srt_mode = self.switch_srt.get() == 1
        diarize_mode = self.switch_diarize.get() == 1
        model_name_ui = self.combo_models.get()

        # 2. VERIFICACIÓN DE MODELO (Igual que antes)
        exists, filename = transcriber.check_model_exists(model_name_ui)
        if not exists:
            msg = f"El modelo '{filename}' no está descargado.\n¿Deseas descargarlo ahora?"
            resp = self.mostrar_confirmacion_oscura("Modelo Faltante", msg)
            if resp:
                # Nota: Si es batch, primero bajamos modelo, luego iniciamos batch
                self.download_and_transcribe(filename, srt_mode, diarize_mode, model_name_ui, is_batch=self.is_batch_mode)
            return

        # 3. MODO BATCH vs NORMAL
        if self.is_batch_mode:
            # Preguntar formato UNA vez
            target_ext = self.ask_export_format()
            if not target_ext: return # Usuario canceló o cerró ventana

            self.prepare_ui_for_process()
            self.btn_process.configure(text="Procesando Cola...")

            # Lanzar Hilo de Cola
            threading.Thread(target=self.run_batch_process,
                           args=(self.queue_files, model_name_ui, srt_mode, diarize_mode, target_ext),
                           daemon=True).start()
        else:
            if not self.selected_file_path: return
            self.prepare_ui_for_process()
            threading.Thread(target=self.run_process, args=(srt_mode, diarize_mode, model_name_ui), daemon=True).start()

    def run_batch_process(self, file_list, model_name, srt_mode, diarize_mode, extension):
        """Procesa la lista de archivos con una barra de progreso GLOBAL basada en el tiempo total."""

        # 1. FASE DE PREPARACIÓN: Calcular duración total de la cola
        self.textbox.delete("0.0", "end")
        self.textbox.insert("0.0", "⏳ Analizando duración total de la cola...\n")

        total_batch_duration = 0
        files_durations = {}

        # Pre-calculamos la duración de cada archivo para ponderar la barra
        for f in file_list:
            duration = transcriber.get_audio_duration(f)
            files_durations[f] = duration
            total_batch_duration += duration

        total_files = len(file_list)
        accumulated_time = 0 # Tiempo acumulado de los archivos ya terminados

        self.textbox.insert("end", f"Total a procesar: {time.strftime('%H:%M:%S', time.gmtime(total_batch_duration))}\n\n")

        # 2. BUCLE DE PROCESAMIENTO
        for index, audio_file in enumerate(file_list):
            if transcriber.is_cancelled: break

            # Datos del archivo actual
            file_name = os.path.basename(audio_file)
            current_file_duration = files_durations.get(audio_file, 0)

            # --- ACTUALIZACIÓN VISUAL DEL HEADER ---
            msg_header = f"--- [{index + 1}/{total_files}] PROCESANDO: {file_name} ---\n"
            msg_header += f"⏱️ Duración: {time.strftime('%M:%S', time.gmtime(current_file_duration))}\n"

            # Escribimos en el textbox sin borrar lo anterior para tener un historial
            # O si prefieres borrar: self.textbox.delete("0.0", "end")
            self.textbox.insert("end", "\n" + msg_header)
            self.textbox.see("end") # Auto-scroll al fondo

            self.title(f"OpenTranscribe (Archivo {index + 1} de {total_files})")

            # --- CALLBACK INTELIGENTE PARA LA BARRA GLOBAL ---
            def batch_progress_callback(local_percent):
                """
                Convierte el % del archivo actual en el % del total de la cola.
                Fórmula: (Tiempo_Acumulado + (Tiempo_Archivo * %_Local)) / Tiempo_Total
                """
                if total_batch_duration > 0:
                    # Cuántos segundos llevamos de ESTE archivo
                    seconds_done_current = current_file_duration * local_percent

                    # Cuántos segundos llevamos EN TOTAL (anteriores + actual)
                    total_seconds_done = accumulated_time + seconds_done_current

                    # Porcentaje global (0.0 a 1.0)
                    global_percent = total_seconds_done / total_batch_duration

                    # Actualizamos la barra
                    self.update_progress(global_percent)

                    # Actualizamos el texto del porcentaje para que sea informativo
                    percent_text = f"{int(global_percent * 100)}% (Total)"
                    self.lbl_progress_percent.configure(text=percent_text)

            # --- EJECUCIÓN ---
            # Variable para acumular texto solo de este archivo para el guardado
            self.current_batch_text_accumulator = ""

            def text_accumulator(text):
                self.current_batch_text_accumulator += text
                # Opcional: Si quieres que salga en pantalla en tiempo real:
                # self.textbox.insert("end", text)
                # self.textbox.see("end")

            transcriber.run_transcription(
                audio_file,
                model_name,
                text_accumulator, # Usamos el acumulador limpio
                batch_progress_callback, # Usamos el nuevo callback global
                with_timestamps=srt_mode,
                diarize=diarize_mode
            )

            # --- AL TERMINAR EL ARCHIVO ---
            if not transcriber.is_cancelled:
                # 1. Guardar
                try:
                    saved_path = self.auto_save_transcript(self.current_batch_text_accumulator, audio_file, extension)
                    self.textbox.insert("end", f"✅ Guardado en: {os.path.basename(saved_path)}\n")
                except Exception as e:
                    self.textbox.insert("end", f"❌ Error guardando: {e}\n")

                # 2. Sumar el tiempo de este archivo al acumulado global
                accumulated_time += current_file_duration
                self.textbox.see("end")

            # Pequeña pausa para respirar
            time.sleep(1)

        # 3. FINALIZACIÓN
        self.after(0, lambda: [
            self.finish_transcription_ui(),
            self.progress_bar.set(1), # Asegurar 100% visual al final
            self.lbl_progress_percent.configure(text="100%"),
            self.mostrar_alerta_oscura("Cola Finalizada", f"Se han procesado {total_files} archivos correctamente.")
        ])

    def prepare_ui_for_process(self):
        self.textbox.delete("0.0", "end")
        self.transcript_segments = []
        self.btn_process.configure(state="disabled", text="Procesando...")
        self.btn_cancel.configure(state="normal")
        self.progress_bar.configure(mode="determinate")
        self.progress_bar.set(0)
        self.unsaved_changes = True
        self.title("OpenTranscribe v2.0 Pro * (Trabajando)")

    def download_and_transcribe(self, filename, srt_mode, diarize_mode, model_name_ui):
        self.prepare_ui_for_process()
        self.btn_process.configure(text="Descargando...")
        self.title("OpenTranscribe (Descargando Modelo...)")

        def thread_target():
            try:
                # 1. Descargar
                self.textbox.insert("end", f"Iniciando descarga de {filename}...\n")
                transcriber.download_model(filename, self.update_progress)
                self.textbox.insert("end", "Descarga completada. Iniciando transcripción...\n\n")

                # 2. Transcribir (Ahora que ya existe)
                # Reiniciamos barra para la transcripción
                self.update_progress(0)
                transcriber.run_transcription(
                    self.selected_file_path,
                    model_name_ui,
                    self.update_text_area,
                    self.update_progress,
                    with_timestamps=srt_mode,
                    diarize=diarize_mode
                )
            except Exception as e:
                self.textbox.insert("end", f"\nError en descarga: {e}")
            finally:
                self.after(0, lambda: [self.finish_transcription_ui(), self.sync_timestamps_from_text()])

        threading.Thread(target=thread_target, daemon=True).start()

    def run_process(self, srt_mode, diarize_mode, model_name):
        # ^^^ FÍJATE AQUÍ: Ahora aceptamos 'diarize_mode' entre los paréntesis

        # Llamamos al backend pasándole el nuevo parámetro
        transcriber.run_transcription(
            self.selected_file_path,
            model_name,
            self.update_text_area,
            self.update_progress,
            with_timestamps=srt_mode,
            diarize=diarize_mode # <--- Se lo pasamos a transcriber.py
        )
        self.after(0, lambda: [self.finish_transcription_ui(), self.sync_timestamps_from_text()])

    def finish_transcription_ui(self):
        self.btn_process.configure(state="normal", text="Transcribir")
        self.btn_cancel.configure(state="disabled")
        if "100%" not in self.lbl_progress_percent.cget("text") and not transcriber.is_cancelled:
             self.progress_bar.set(1)
             self.lbl_progress_percent.configure(text="100%")

    def parse_and_insert_line(self, text_line):
        # Limpiamos posibles espacios extra
        clean_line = text_line

        # Detectamos patrón de hablante [SPEAKER_00] o (SPEAKER_00)
        # Regex busca: (cualquier cosa antes)(SPEAKER_XX)(resto del texto)
        pattern = re.compile(r"(.*)(\[SPEAKER_\d+\]|\(SPEAKER_\d+\))(.*)")
        match = pattern.search(clean_line)

        if match:
            # Si encontramos hablante, lo formateamos bonito
            prefix = match.group(1) # Tiempos si los hay
            speaker_tag = match.group(2) # El tag feo
            content = match.group(3) # Lo que dicen

            # Extraemos solo el número
            num = "".join(filter(str.isdigit, speaker_tag))
            formatted_speaker = f" 👤 Hablante {int(num) + 1}: " # +1 para que empiece en 1, no en 0

            # Insertamos el prefijo (tiempos) normal
            self.textbox.insert("end", prefix)

            # Insertamos el hablante con COLOR y NEGRITA (simulada con tags)
            self.textbox.tag_config(f"speaker_{num}", foreground="#4da6ff", font=("Roboto", 14, "bold")) # Azul claro
            self.textbox.insert("end", formatted_speaker, f"speaker_{num}")

            # Insertamos el contenido normal
            self.textbox.insert("end", content + "\n")
        else:
            # Si no hay hablante, insertamos normal
            self.textbox.insert("end", clean_line)

        self.textbox.see("end")

    def sync_timestamps_from_text(self):
        self.transcript_segments = []
        total_lines = int(self.textbox.index("end-1c").split('.')[0])
        pattern = re.compile(r"\[(\d{2}):(\d{2}):(\d{2}\.\d{3}) --> (\d{2}):(\d{2}):(\d{2}\.\d{3})\]")
        for i in range(1, total_lines + 1):
            line_idx = f"{i}.0"
            line_end_idx = f"{i}.end"
            line_text = self.textbox.get(line_idx, line_end_idx)
            match = pattern.search(line_text)
            if match:
                try:
                    h1, m1, s1 = map(float, match.group(1, 2, 3))
                    start_seconds = h1 * 3600 + m1 * 60 + s1
                    h2, m2, s2 = map(float, match.group(4, 5, 6))
                    end_seconds = h2 * 3600 + m2 * 60 + s2
                    segment_data = {'start': start_seconds, 'end': end_seconds, 'idx_start': line_idx, 'idx_end': line_end_idx}
                    self.transcript_segments.append(segment_data)
                except ValueError: continue

    def sync_timestamps_manual(self):
        self.sync_timestamps_from_text()
        self.mostrar_alerta_oscura("Sincronizado", "Se han actualizado los tiempos del Karaoke.")

    def update_audio_slider_loop(self):
        if self.is_playing and self.total_duration > 0:
            current_time = self.current_offset + (pygame.mixer.music.get_pos() / 1000)
            if current_time > self.total_duration:
                self.stop_audio()
                return
            self.slider_audio.set(current_time / self.total_duration)
            current_str = time.strftime('%M:%S', time.gmtime(int(current_time)))
            total_str = time.strftime('%M:%S', time.gmtime(self.total_duration))
            self.lbl_audio_time.configure(text=f"{current_str} / {total_str}")
            for seg in self.transcript_segments:
                if seg['start'] <= current_time <= seg['end']:
                    self.textbox._textbox.tag_add("highlight", seg['idx_start'], seg['idx_end'])
                else:
                    self.textbox._textbox.tag_remove("highlight", seg['idx_start'], seg['idx_end'])
            self.after(100, self.update_audio_slider_loop)

    def toggle_audio(self):
        if not self.is_playing:
            try:
                if len(self.transcript_segments) == 0: self.sync_timestamps_from_text()
                if pygame.mixer.music.get_pos() == -1: pygame.mixer.music.play()
                else:
                     pygame.mixer.music.unpause()
                     if not pygame.mixer.music.get_busy(): pygame.mixer.music.play(start=self.current_offset)
                self.is_playing = True
                self.btn_play.configure(text="⏸ Pausa", fg_color="#555")
                self.update_audio_slider_loop()
            except Exception as e: print(e)
        else:
            pygame.mixer.music.pause()
            self.is_playing = False
            self.btn_play.configure(text="▶ Reanudar", fg_color="#333")

    def select_file(self):
        """Abre el explorador nativo del sistema."""
        # Filtros de archivo
        file_types = [
            ("Todos los medios", "*.mp3 *.wav *.m4a *.mp4 *.mkv *.mov *.avi *.webm *.flv"),
            ("Audio", "*.mp3 *.wav *.m4a"),
            ("Vídeo", "*.mp4 *.mkv *.mov *.avi *.webm *.flv"),
            ("Todos los archivos", "*.*")
        ]

        # Usamos filedialog nativo (permite selección múltiple)
        filenames = filedialog.askopenfilenames(
            title="Seleccionar Archivos",
            filetypes=file_types
        )

        if not filenames:
            return # Cancelado

        # Convertimos a lista
        filepaths = list(filenames)

        if len(filepaths) == 1:
            self.cargar_archivo_comun(filepaths[0])
        else:
            # Lógica de Cola
            self.al_soltar_archivo(type('Event', (object,), {'data': " ".join(filepaths)})())

    def al_soltar_archivo(self, event):
        filepaths = self.parse_dropped_files(event.data)

        if not filepaths:
            return

        valid_exts = ['.mp3', '.wav', '.m4a', '.mp4', '.mkv', '.mov', '.avi', '.webm', '.flv']

        if len(filepaths) == 1:
            # Comportamiento normal (1 archivo)
            self.queue_files = []
            self.is_batch_mode = False
            self.cargar_archivo_comun(filepaths[0])
        else:
            # MODO BATCH
            self.queue_files = filepaths
            self.is_batch_mode = True

            # Validar extensiones
            valid_files = [f for f in filepaths if any(f.lower().endswith(ext) for ext in valid_exts)]
            self.queue_files = valid_files

            if not valid_files:
                self.mostrar_alerta_oscura("Error", "Ningún archivo válido detectado.")
                return

            # Actualizar UI
            self.lbl_filename.configure(text=f"📚 COLA: {len(valid_files)} archivos listos", text_color="#4da6ff")
            self.btn_process.configure(state="normal", text="Procesar Cola 📚")
            self.btn_play.configure(state="disabled") # No reproducimos en modo cola
            self.textbox.delete("0.0", "end")
            self.textbox.insert("0.0", "Modo Cola activado.\nArchivos detectados:\n\n")
            for f in valid_files:
                self.textbox.insert("end", f"• {os.path.basename(f)}\n")

    def cargar_archivo_comun(self, filename):
        # 1. Validar extensión
        extensiones_validas = ['.mp3', '.wav', '.m4a', '.mp4', '.mkv', '.mov', '.avi', '.webm', '.flv']

        if not any(filename.lower().endswith(ext) for ext in extensiones_validas):
            self.mostrar_alerta_oscura("Error", "Formato no soportado.")
            return

        # 2. Resetear variables de Cola (IMPORTANTE para salir del modo Batch)
        self.is_batch_mode = False
        self.queue_files = []

        # 3. Actualizar Referencias
        self.selected_file_path = filename
        self.lbl_filename.configure(text=os.path.basename(filename), text_color="white")

        # 4. RESTAURAR UI (Aquí estaba el fallo)
        # Volvemos el botón a su texto normal
        self.btn_process.configure(state="normal", text="TRANSCRIBIR")

        # Limpiamos la caja de texto (borramos la lista de la cola anterior)
        self.textbox.delete("0.0", "end")
        self.textbox.insert("0.0", "El texto transcrito aparecerá aquí...\n")

        # Reseteamos barras de progreso
        self.progress_bar.set(0)
        self.lbl_progress_percent.configure(text="0%")

        # 5. Cargar Previsualización de Audio
        self.load_audio_preview()

    def load_audio_preview(self):
        if self.selected_file_path:
            try:
                self.total_duration = transcriber.get_audio_duration(self.selected_file_path)
                pygame.mixer.music.load(self.selected_file_path)
                self.slider_audio.set(0)
                self.current_offset = 0
                self.is_playing = False
                self.btn_play.configure(state="normal", text="▶ Reproducir", fg_color="#333")
                self.btn_stop.configure(state="normal")
                total_str = time.strftime('%M:%S', time.gmtime(self.total_duration))
                self.lbl_audio_time.configure(text=f"00:00 / {total_str}")
                self.sync_timestamps_from_text()
            except Exception as e:
                print(f"Error cargando audio: {e}")
                self.btn_play.configure(state="disabled")

    def stop_audio(self):
        pygame.mixer.music.stop()
        self.is_playing = False
        self.current_offset = 0
        self.slider_audio.set(0)
        self.textbox._textbox.tag_remove("highlight", "1.0", "end")
        self.btn_play.configure(text="▶ Reproducir", fg_color="#333")
        total_str = time.strftime('%M:%S', time.gmtime(self.total_duration))
        self.lbl_audio_time.configure(text=f"00:00 / {total_str}")

    def seek_audio(self, value):
        if self.total_duration > 0:
            target_time = value * self.total_duration
            self.current_offset = target_time
            pygame.mixer.music.play(start=target_time)
            if not self.is_playing: pygame.mixer.music.pause()
            else:
                self.is_playing = True
                self.btn_play.configure(text="⏸ Pausa", fg_color="#555")
                self.update_audio_slider_loop()

    def update_text_area(self, text):
        self.parse_and_insert_line(text)

    def update_progress(self, progress_float):
        self.after(0, lambda: self._update_progress_gui(progress_float))

    def _update_progress_gui(self, value):
        self.progress_bar.set(value)
        percent_text = f"{int(value * 100)}%"
        self.lbl_progress_percent.configure(text=percent_text)

    def save_text(self):
        """Guardado nativo robusto."""
        if self.is_batch_mode:
            self.mostrar_alerta_oscura("Aviso", "En modo Cola el guardado es automático.")
            return

        raw_content = self.textbox.get("0.0", "end").strip()
        if not raw_content or "El texto transcrito aparecerá aquí" in raw_content:
            self.mostrar_alerta_oscura("Error", "No hay texto para guardar.")
            return

        # Filtros
        filtros = [
            ("Documento Word (*.docx)", "*.docx") if HAS_DOCX else None,
            ("Texto Plano (*.txt)", "*.txt"),
            ("Subtítulos SRT (*.srt)", "*.srt"),
            ("Subtítulos VTT (*.vtt)", "*.vtt"),
            ("Excel / CSV (*.csv)", "*.csv")
        ]
        filtros = [f for f in filtros if f is not None]

        # Diálogo nativo
        filename = filedialog.asksaveasfilename(
            title="Guardar Transcripción",
            defaultextension=".docx" if HAS_DOCX else ".txt",
            filetypes=filtros
        )

        if not filename:
            return

        ext = os.path.splitext(filename)[1].lower()

        try:
            # 3. PARSEO DEL TEXTO (Convertir texto plano a estructura de datos)
            lines_data = []

            # Regex para capturar tiempos: [00:00:00.000 --> 00:00:05.000]
            timestamp_pattern = re.compile(r"\[(\d{2}:\d{2}:\d{2}[\.,]\d{3}) --> (\d{2}:\d{2}:\d{2}[\.,]\d{3})\]")

            raw_lines = raw_content.split('\n')
            current_seg = {"start": "", "end": "", "text": ""}

            for line in raw_lines:
                match = timestamp_pattern.search(line)
                if match:
                    # Guardar segmento previo
                    if current_seg["start"]:
                        lines_data.append(current_seg)

                    # Normalizar tiempos (usar punto internamente)
                    s_time = match.group(1).replace(',', '.')
                    e_time = match.group(2).replace(',', '.')

                    # Limpiar el texto de la marca de tiempo
                    txt = timestamp_pattern.sub("", line).strip()

                    current_seg = {"start": s_time, "end": e_time, "text": txt}
                else:
                    # Texto continuado (sin tiempo)
                    if line.strip():
                        if current_seg["start"]:
                            current_seg["text"] += " " + line.strip()
                        else:
                            # Texto huérfano (headers, notas)
                            lines_data.append({"start": "", "end": "", "text": line.strip()})

            # Añadir el último segmento
            if current_seg["start"] or current_seg["text"]:
                lines_data.append(current_seg)

            # 4. ESCRITURA SEGÚN EL FORMATO

            # --- A) MICROSOFT WORD (.docx) ---
            if ext == ".docx" and HAS_DOCX:
                doc = Document()
                doc.add_heading('Transcripción - OpenTranscribe', 0)

                for item in lines_data:
                    p = doc.add_paragraph()

                    # Tiempo en Azul
                    if item["start"]:
                        run_time = p.add_run(f"[{item['start']} - {item['end']}] ")
                        run_time.bold = True
                        run_time.font.color.rgb = RGBColor(0, 50, 150)

                    # Detección de Hablantes (Rojo)
                    text_content = item["text"]
                    if "👤" in text_content:
                        parts = text_content.split(":", 1)
                        if len(parts) > 1:
                            run_speaker = p.add_run(parts[0] + ":")
                            run_speaker.bold = True
                            run_speaker.font.color.rgb = RGBColor(200, 0, 0)
                            p.add_run(parts[1])
                        else:
                            p.add_run(text_content)
                    else:
                        p.add_run(text_content)

                doc.save(filename)

            # --- B) EXCEL / CSV (.csv) ---
            elif ext == ".csv":
                with open(filename, mode='w', newline='', encoding='utf-8-sig') as csv_file:
                    writer = csv.writer(csv_file, delimiter=';')
                    writer.writerow(['Inicio', 'Fin', 'Contenido'])
                    for item in lines_data:
                        if item["start"]:
                            writer.writerow([item["start"], item["end"], item["text"]])

            # --- C) SUBTÍTULOS VTT (.vtt) ---
            elif ext == ".vtt":
                with open(filename, "w", encoding="utf-8") as f:
                    f.write("WEBVTT\n\n")
                    counter = 1
                    for item in lines_data:
                        if item["start"]:
                            f.write(f"{counter}\n")
                            f.write(f"{item['start']} --> {item['end']}\n")
                            f.write(f"{item['text']}\n\n")
                            counter += 1

            # --- D) SUBTÍTULOS SRT (.srt) ---
            elif ext == ".srt":
                with open(filename, "w", encoding="utf-8") as f:
                    counter = 1
                    for item in lines_data:
                        if item["start"]:
                            f.write(f"{counter}\n")
                            # SRT requiere coma en milisegundos
                            start_srt = item['start'].replace('.', ',')
                            end_srt = item['end'].replace('.', ',')
                            f.write(f"{start_srt} --> {end_srt}\n")
                            f.write(f"{item['text']}\n\n")
                            counter += 1

            # --- E) TEXTO PLANO (.txt) ---
            else:
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(raw_content)

            # 5. FINALIZACIÓN
            self.unsaved_changes = False
            self.title("OpenTranscribe v2.0 Pro")
            self.mostrar_alerta_oscura("Guardado", f"Archivo guardado exitosamente:\n{os.path.basename(filename)}")

        except Exception as e:
            print(f"Error guardando: {e}")
            self.mostrar_alerta_oscura("Error", f"No se pudo guardar el archivo:\n{str(e)}")

    def open_find_replace_dialog(self):
        self.dialog = ctk.CTkToplevel(self)
        self.dialog.title("Buscar y Reemplazar")
        self.dialog.geometry("400x250")
        self.dialog.attributes("-topmost", True)
        ctk.CTkLabel(self.dialog, text="Buscar palabra:").pack(pady=(20, 5))
        entry_find = ctk.CTkEntry(self.dialog, width=250)
        entry_find.pack(pady=5)
        ctk.CTkLabel(self.dialog, text="Reemplazar con:").pack(pady=(10, 5))
        entry_replace = ctk.CTkEntry(self.dialog, width=250)
        entry_replace.pack(pady=5)
        def ejecutar_reemplazo():
            texto_a_buscar = entry_find.get()
            texto_nuevo = entry_replace.get()
            if texto_a_buscar: self.perform_replace(texto_a_buscar, texto_nuevo, parent_window=self.dialog)
        ctk.CTkButton(self.dialog, text="Reemplazar Todo", command=ejecutar_reemplazo, fg_color="#1f6aa5").pack(pady=20)

    def perform_replace(self, old_text, new_text, parent_window):
        content = self.textbox.get("0.0", "end")
        if old_text not in content:
            self.mostrar_alerta_oscura("Error", f"No se encontró '{old_text}'", parent_window)
            return
        new_content = content.replace(old_text, new_text)
        self.textbox.delete("0.0", "end")
        self.textbox.insert("0.0", new_content)
        self.mark_as_modified()
        self.sync_timestamps_from_text()
        self.mostrar_alerta_oscura("Éxito", f"Se reemplazó '{old_text}' por '{new_text}'.", parent_window)

    def mostrar_alerta_oscura(self, titulo, mensaje, parent_window=None):
        padre = parent_window if parent_window else self
        popup = ctk.CTkToplevel(padre)
        popup.title(titulo)
        popup.geometry("350x180")
        popup.resizable(False, False)
        popup.grid_columnconfigure(0, weight=1)
        popup.grid_rowconfigure(0, weight=1)
        lbl = ctk.CTkLabel(popup, text=mensaje, font=("Roboto", 14), wraplength=300)
        lbl.pack(pady=40, padx=20)
        btn = ctk.CTkButton(popup, text="Aceptar", command=popup.destroy, width=100)
        btn.pack(pady=(0, 20))
        popup.attributes("-topmost", True)
        popup.update()
        try: popup.grab_set()
        except: pass
        padre.wait_window(popup)

    def mostrar_confirmacion_oscura(self, titulo, mensaje):
        dialog = ctk.CTkToplevel(self)
        dialog.title(titulo)
        dialog.geometry("350x180")
        dialog.resizable(False, False)
        ctk.CTkLabel(dialog, text=mensaje, font=("Roboto", 14), wraplength=300).pack(pady=30, padx=20)
        frame_btns = ctk.CTkFrame(dialog, fg_color="transparent")
        frame_btns.pack(pady=10)
        self.respuesta_usuario = False
        def on_yes():
            self.respuesta_usuario = True
            dialog.destroy()
        def on_no():
            self.respuesta_usuario = False
            dialog.destroy()
        ctk.CTkButton(frame_btns, text="Sí, salir", command=on_yes, fg_color="red", hover_color="#800000", width=100).pack(side="left", padx=10)
        ctk.CTkButton(frame_btns, text="Cancelar", command=on_no, fg_color="gray", width=100).pack(side="right", padx=10)
        dialog.attributes("-topmost", True)
        dialog.update()
        try: dialog.grab_set()
        except: pass
        self.wait_window(dialog)
        return self.respuesta_usuario

    def mark_as_modified(self, event=None):
        keys_to_ignore = ["Control_L", "Control_R", "Alt_L", "Shift_L", "Up", "Down", "Left", "Right"]
        if event and event.keysym in keys_to_ignore: return
        if not self.unsaved_changes:
            self.unsaved_changes = True
            self.title("OpenTranscribe * (Sin guardar)")

    def on_closing(self):
        if self.unsaved_changes:
            # Usamos nuestra nueva alerta oscura
            salir = self.mostrar_confirmacion_oscura("Salir", "Tienes cambios sin guardar.\n¿Estás seguro de que quieres salir?")
            if salir:
                try: pygame.mixer.quit()
                except: pass
                self.destroy()
        else:
            try: pygame.mixer.quit()
            except: pass
            self.destroy()

    def abrir_ayuda(self):
        # Crear ventana emergente más grande
        ayuda_window = ctk.CTkToplevel(self)
        ayuda_window.title("Manual de Usuario - OpenTranscribe")
        ayuda_window.geometry("550x750") # Más alto para que quepa todo
        ayuda_window.resizable(False, True) # Permitir redimensionar alto
        ayuda_window.attributes("-topmost", True)

        # 1. LOGO (Depurado y Ajustado)
        try:
            logo_path = "Logo.jpg"
            # Soporte para modo congelado (exe/binario)
            if hasattr(sys, '_MEIPASS'):
                logo_path = os.path.join(sys._MEIPASS, "Logo.jpg")
            # Soporte para modo desarrollo local
            elif not os.path.exists(logo_path):
                logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Logo.jpg")

            if os.path.exists(logo_path):
                img_data = Image.open(logo_path)

                # --- CONFIGURACIÓN DE TAMAÑO ---
                # La ventana mide 550px de ancho.
                # 350px es un tamaño muy bueno (grande pero cabe con márgenes).
                ancho_deseado = 350

                # Calculamos el alto proporcional (Regla de tres)
                w_original, h_original = img_data.size
                ratio = ancho_deseado / float(w_original)
                alto_calculado = int(float(h_original) * float(ratio))

                print(f"DEBUG: Cargando logo. Original: {w_original}x{h_original} -> Nuevo: {ancho_deseado}x{alto_calculado}")

                logo_img = ctk.CTkImage(light_image=img_data, dark_image=img_data,
                                      size=(ancho_deseado, alto_calculado))

                lbl_img = ctk.CTkLabel(ayuda_window, text="", image=logo_img)
                lbl_img.pack(pady=(20, 10))
            else:
                print(f"DEBUG: No se encontró el archivo en: {logo_path}")

        except Exception as e:
            print(f"ERROR cargando el logo: {e}")

        ctk.CTkLabel(ayuda_window, text="OpenTranscribe v2.0", font=("Roboto Medium", 20)).pack(pady=5)

        # 2. ÁREA DE TEXTO CON SCROLL (Para todo el manual)
        # Usamos Textbox en modo lectura para que sea scrollable y copiable
        info_text = ctk.CTkTextbox(ayuda_window, width=500, height=480, corner_radius=10,
                                   fg_color="#232323", text_color="#eeeeee", font=("Consolas", 12))
        info_text.pack(pady=10, padx=20, fill="both", expand=True)

        # --- CONTENIDO DEL MANUAL ---
        manual = (
            "============================================\n"
            "GUÍA DE FUNCIONES PRINCIPALES\n"
            "============================================\n\n"

            "1. TRANCRIPCIÓN BÁSICA\n"
            "----------------------\n"
            "• Arrastra un archivo de audio o vídeo a la ventana.\n"
            "• Elige el 'Modelo IA' (Base es recomendado).\n"
            "• Pulsa 'TRANSCRIBIR'.\n\n"

            "2. MODO COLA (Lotes / Batch) [NUEVO] 📚\n"
            "------------------------------------\n"
            "• Arrastra MÚLTIPLES archivos a la vez (ej. 10 vídeos).\n"
            "• La aplicación detectará el modo 'Cola'.\n"
            "• Pulsa 'Procesar Cola'.\n"
            "• El sistema te preguntará el formato (Word, PDF, etc.).\n"
            "• Los archivos se guardarán AUTOMÁTICAMENTE en la\n"
            "  misma carpeta que los originales.\n\n"

            "3. SOPORTE MULTIMEDIA 🎬\n"
            "------------------------\n"
            "• Aceptamos: MP3, WAV, M4A (Audio).\n"
            "• Aceptamos: MP4, MKV, AVI, MOV, WEBM (Vídeo).\n"
            "• El vídeo se procesa internamente, no necesitas\n"
            "  extraer el audio antes.\n\n"

            "4. HERRAMIENTAS INTELIGENTES\n"
            "----------------------------\n"
            "• Detectar Hablantes 👥: Intenta distinguir quién habla\n"
            "  (Hablante 1, Hablante 2...).\n"
            "• Modo Subtítulos: Genera marcas de tiempo exactas\n"
            "  para crear archivos .SRT o .VTT.\n"
            "• Reproductor Karaoke: Pulsa ▶ para escuchar el audio\n"
            "  y ver cómo se resalta el texto en tiempo real.\n"
            "• Sincronizar: Si editas el texto manualmente, pulsa\n"
            "  este botón para recalcular los tiempos del karaoke.\n\n"

            "5. EXPORTACIÓN\n"
            "--------------\n"
            "• Word (.docx): Con colores y negritas.\n"
            "• Subtítulos (.srt/.vtt): Listos para YouTube/VLC.\n"
            "• Excel (.csv): Para análisis de datos.\n"
            "• Texto (.txt): Simple y ligero.\n"
        )

        info_text.insert("0.0", manual)
        info_text.configure(state="disabled") # Hacemos que sea solo lectura

        # 3. SECCIÓN DE CONTACTO
        frame_contact = ctk.CTkFrame(ayuda_window, fg_color="transparent")
        frame_contact.pack(pady=10, fill="x")

        ctk.CTkLabel(frame_contact, text="¿Dudas o Bugs?", font=("Roboto", 12, "bold")).pack()

        # Email Clicable
        lbl_mail = ctk.CTkLabel(frame_contact, text="anabasasoft@gmail.com", text_color="#4da6ff", cursor="hand2")
        lbl_mail.pack()
        lbl_mail.bind("<Button-1>", lambda e: webbrowser.open("mailto:anabasasoft@gmail.com"))

        # Web Clicable
        lbl_web = ctk.CTkLabel(frame_contact, text="anabasasoft.github.io", text_color="#4da6ff", cursor="hand2")
        lbl_web.pack()
        lbl_web.bind("<Button-1>", lambda e: webbrowser.open("https://anabasasoft.github.io"))

        # Botón Cerrar
        ctk.CTkButton(ayuda_window, text="Cerrar", command=ayuda_window.destroy,
                      fg_color="#333", hover_color="#444", width=100).pack(pady=(5, 20))

    def check_system_requirements(self):
        """Verifica dependencias del sistema"""
        missing = []

        # 1. Verificar FFMPEG (Esto depende del usuario)
        if not transcriber.get_ffmpeg_executable():
            missing.append("FFmpeg (Necesario para procesar audio)")

        # 2. Verificar Whisper (Debería estar incluido)
        if not transcriber.get_whisper_executable():
            missing.append("ERROR: No se encuentra el motor interno (whisper-cli)")

        if missing:
            msg = "Faltan componentes necesarios:\n\n"
            for item in missing:
                msg += f"❌ {item}\n"

            if shutil.which("apt"): # Si es Debian/Ubuntu/Mint
                msg += "\nIntenta instalar FFmpeg con:\nsudo apt install ffmpeg"

            self.mostrar_alerta_oscura("Faltan Dependencias", msg)

    def parse_dropped_files(self, data):
        """Convierte el string de TkinterDnD en una lista de rutas limpias."""
        # Si viene entre corchetes {} (común en Linux/Windows con espacios)
        # Usamos regex para separar
        files = []
        if data.startswith('{') or '}' in data:
            parts = re.findall(r'\{.*?\}|\S+', data)
            for part in parts:
                path = part.strip('{}')
                if os.path.isfile(path):
                    files.append(path)
        else:
            # Caso simple: un archivo o varios sin espacios
            candidates = data.split()
            for c in candidates:
                if os.path.isfile(c):
                    files.append(c)

        # Fallback: si el regex falla, intentamos usar data directo si es un archivo
        if not files and os.path.isfile(data):
            files = [data]

        return files

    def ask_export_format(self):
        """Pregunta al usuario en qué formato guardar los archivos de la cola."""
        dialog = ctk.CTkToplevel(self)
        dialog.title("Formato de Salida")
        dialog.geometry("300x250")
        dialog.resizable(False, False)
        dialog.attributes("-topmost", True)

        self.selected_format = None

        ctk.CTkLabel(dialog, text="Elige el formato para guardar:", font=("Roboto", 14, "bold")).pack(pady=20)

        def set_format(fmt):
            self.selected_format = fmt
            dialog.destroy()

        ctk.CTkButton(dialog, text="📄 Word (.docx)", command=lambda: set_format(".docx") if HAS_DOCX else None,
                      state="normal" if HAS_DOCX else "disabled", fg_color="#2b5797").pack(pady=5)
        ctk.CTkButton(dialog, text="📝 Texto (.txt)", command=lambda: set_format(".txt"), fg_color="#444").pack(pady=5)
        ctk.CTkButton(dialog, text="🎬 Subtítulos (.srt)", command=lambda: set_format(".srt"), fg_color="#d68a00").pack(pady=5)
        ctk.CTkButton(dialog, text="📊 Excel/CSV (.csv)", command=lambda: set_format(".csv"), fg_color="#217346").pack(pady=5)

        dialog.wait_window(dialog)
        return self.selected_format

    def auto_save_transcript(self, text_content, audio_path, extension):
        """
        Guarda la transcripción automáticamente en la misma carpeta del audio original,
        soportando todos los formatos (.docx, .txt, .srt, .vtt, .csv).
        """
        if not text_content.strip():
            return None

        # 1. Generar nombre de archivo basado en el audio original
        # Ejemplo: /home/user/Entrevista.mp3 -> /home/user/Entrevista_Transcribed.docx
        folder = os.path.dirname(audio_path)
        base_name = os.path.splitext(os.path.basename(audio_path))[0]
        filename = os.path.join(folder, f"{base_name}_Transcribed{extension}")

        try:
            # --- FASE A: PARSEO DEL TEXTO (Convertir texto plano a datos estructurados) ---
            lines_data = []
            raw_lines = text_content.split('\n')
            # Regex para detectar tiempos: [00:00:00.000 --> 00:00:05.000]
            timestamp_pattern = re.compile(r"\[(\d{2}:\d{2}:\d{2}[\.,]\d{3}) --> (\d{2}:\d{2}:\d{2}[\.,]\d{3})\]")

            current_segment = {"start": "", "end": "", "text": ""}

            for line in raw_lines:
                match = timestamp_pattern.search(line)
                if match:
                    # Si encontramos una línea de tiempo, guardamos el segmento anterior y empezamos uno nuevo
                    if current_segment["start"]:
                        lines_data.append(current_segment)

                    start_time = match.group(1).replace(',', '.')
                    end_time = match.group(2).replace(',', '.')

                    # Limpiamos el texto quitando la marca de tiempo
                    text_part = timestamp_pattern.sub("", line).strip()

                    current_segment = {"start": start_time, "end": end_time, "text": text_part}
                else:
                    # Si es una línea de texto continuado sin tiempo (o el header)
                    if line.strip():
                        if current_segment["start"]: # Solo añadimos si estamos dentro de un segmento válido
                            current_segment["text"] += " " + line.strip()
                        # Nota: Ignoramos el header de "Procesando archivo..." si no tiene timestamp

            # Añadir el último segmento que quedó pendiente
            if current_segment["start"]:
                lines_data.append(current_segment)


            # --- FASE B: GUARDADO SEGÚN EXTENSIÓN ---

            # 1. WORD (.docx)
            if extension == ".docx" and HAS_DOCX:
                doc = Document()
                doc.add_heading(f'Transcripción: {base_name}', 0)

                for item in lines_data:
                    p = doc.add_paragraph()

                    # Estilo del tiempo (Azul)
                    run_time = p.add_run(f"[{item['start']} - {item['end']}] ")
                    run_time.bold = True
                    run_time.font.color.rgb = RGBColor(0, 50, 150)

                    # Estilo del texto (Detectando hablantes si los hay)
                    text_content_seg = item["text"]
                    if "👤" in text_content_seg:
                        parts = text_content_seg.split(":", 1)
                        if len(parts) > 1:
                            run_speaker = p.add_run(parts[0] + ":")
                            run_speaker.bold = True
                            run_speaker.font.color.rgb = RGBColor(200, 0, 0) # Rojo oscuro
                            p.add_run(parts[1])
                        else:
                            p.add_run(text_content_seg)
                    else:
                        p.add_run(text_content_seg)

                doc.save(filename)

            # 2. EXCEL / CSV (.csv)
            elif extension == ".csv":
                with open(filename, mode='w', newline='', encoding='utf-8-sig') as csv_file:
                    writer = csv.writer(csv_file, delimiter=';')
                    writer.writerow(['Inicio', 'Fin', 'Contenido'])
                    for item in lines_data:
                        writer.writerow([item["start"], item["end"], item["text"]])

            # 3. SUBTÍTULOS VTT (.vtt)
            elif extension == ".vtt":
                with open(filename, "w", encoding="utf-8") as f:
                    f.write("WEBVTT\n\n")
                    for i, item in enumerate(lines_data):
                        f.write(f"{i+1}\n")
                        # VTT usa puntos para milisegundos (00:00:00.000)
                        f.write(f"{item['start']} --> {item['end']}\n")
                        f.write(f"{item['text']}\n\n")

            # 4. SUBTÍTULOS SRT (.srt)
            elif extension == ".srt":
                with open(filename, "w", encoding="utf-8") as f:
                    for i, item in enumerate(lines_data):
                        f.write(f"{i+1}\n")
                        # SRT usa comas para milisegundos (00:00:00,000)
                        start_srt = item['start'].replace('.', ',')
                        end_srt = item['end'].replace('.', ',')
                        f.write(f"{start_srt} --> {end_srt}\n")
                        f.write(f"{item['text']}\n\n")

            # 5. TEXTO PLANO (.txt)
            else:
                # Para TXT guardamos todo el contenido raw (incluyendo headers si los hubiera)
                # o reconstruimos limpio. Aquí guardamos el raw original para mantener consistencia.
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(text_content)

            return filename

        except Exception as e:
            print(f"Error guardando automático ({extension}): {e}")
            return None

if __name__ == "__main__":
    app = OpenTranscribeApp()
    app.mainloop()
